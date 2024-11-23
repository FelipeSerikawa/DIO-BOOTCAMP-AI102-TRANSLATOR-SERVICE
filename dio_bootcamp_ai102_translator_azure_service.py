#!pip install python-docx

#TRADUTOR UTILIZANDO O SERVIÇO DE TRADUÇÃO NA NUVEM DO AZURE
import requests
from docx import Document
import os

#INSERIR A CHAVE DO SERVIÇO CRIADO
subscription_key = ""
endpoint = "https://api.cognitive.microsofttranslator.com/"
location = "brazilsouth"

def translator_text(from_language, to_language, text):
  path = "/translate"
  constructed_url = endpoint + path
  headers = {
      "Ocp-Apim-Subscription-Key": subscription_key,
      "Ocp-Apim-Subscription-Region": location,
      "Content-type": "application/json",
      "X-ClientTraceId": str(os.urandom(16))
  }

  body = [{
      "text": text
  }]

  params = {
      'api-version': '3.0',
      'from': [from_language],
      'to': [to_language]
  }
  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

text = "Hi!"
from_language = "en"
to_language = "pt-br"
translator_text(from_language, to_language, text)

def translator_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text("en", "pt-br", paragraph.text)
    full_text.append(translated_text)
  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)
  path_translated = path.replace(".docx", "_translated.docx")
  translated_doc.save(path_translated)
  return path_translated

input_file = "/content/teste.docx"
translator_document(input_file)